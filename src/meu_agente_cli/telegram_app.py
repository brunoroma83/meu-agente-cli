import sys
import logging
import os
import re
import time
import subprocess
import telebot
from rich.console import Console
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton

# Adiciona o diretório raiz ao path para importação
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from meu_agente_cli import config, db, agent
from meu_agente_cli.main import handle_slash_command, initialize_components

# Carrega credenciais do Telegram (não requer banco de dados durante a importação)
token = config.get_telegram_bot_token()
authorized_ids = config.get_telegram_authorized_user_ids()

if not token:
    if __name__ == "__main__":
        print("[ERRO] Token do Telegram não encontrado! Configure a variável TELEGRAM_BOT_TOKEN no .env ou config.json.")
        sys.exit(1)
    token = "123456:dummy_token_for_import_validation"

bot = telebot.TeleBot(token)

pending_transcriptions = {}  # {chat_id: {"text": str, "audio_path": str}}
user_states = {}  # {chat_id: str}

TELEGRAM_INSTRUCTION = (
    "\n\n[INSTRUÇÃO DO SISTEMA: Você está respondendo via Telegram. Se for responder diretamente "
    "ao usuário em texto, formate com listas limpas, quebras de linhas duplas e emojis, sem tabelas ASCII "
    "ou caixas unicode complexas, pois serão exibidas em uma tela estreita de celular. Se você precisar "
    "acionar uma ferramenta, continue respondendo APENAS com o bloco JSON da ferramenta, sem texto adicional de conversa.]"
)


def is_authorized(message) -> bool:
    """Verifica se o chat ID remetente está na lista de usuários autorizados."""
    if not authorized_ids:
        # Por padrão de segurança reforçada, se a lista estiver vazia, bloqueia tudo
        return False
    return message.chat.id in authorized_ids

def is_logged_in(message) -> bool:
    """Verifica se há um usuário logado e ativo no banco."""
    logged_user = db.get_logged_in_user()
    if not logged_user:
        bot.reply_to(message, "⚠️ Acesso negado. Nenhum usuário logado. Por favor, faça login usando `/login <nome_usuario>`.")
        return False
    return True

def strip_ansi(text: str) -> str:
    """Remove códigos de cores ANSI do console do Rich."""
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    return ansi_escape.sub('', text)

def format_help_telegram() -> str:
    return (
        "🛡️ *Meu Agente CLI - Comandos Disponíveis:*\n\n"
        "• `/help` - Mostra esta lista de ajuda.\n"
        "• `/login <nome_usuario>` - Inicia uma sessão de login (válida por 24h).\n"
        "• `/logout` - Encerra a sessão ativa.\n"
        "• `/status` - Mostra conexões e estado atual de segurança.\n"
        "• `/clear` - Limpa o histórico de conversa com o agente.\n"
        "• `/history <limite>` - Exibe ou altera o tamanho do histórico.\n"
        "• `/models` - Lista e altera os modelos de LLM (locais ou externos).\n"
        "• `/safe` / `/unsafe` - Ativa ou desativa o Modo Seguro.\n"
        "• `/notes` - Lista todas as anotações salvas.\n"
        "• `/finance` - Resumo financeiro do mês atual.\n"
        "  └─ _Filtros:_ `/finance next`, `/finance all`, `/finance deleted`, `/finance mes=MM-YYYY`, `/finance q=busca`.\n"
        "• `/finance card list` - Lista cartões de crédito cadastrados.\n"
        "• `/finance card add <nome> <fechamento> <vencimento>` - Cadastra cartão.\n"
        "• `/finance card buy <cartao> <categoria> <valor> <parcelas> <descricao> [data]` - Lança despesa no cartão.\n"
        "• `/cron` - Lista tarefas agendadas em segundo plano.\n"
        "• `/reunioes` - Lista reuniões e transcrições de áudio salvas.\n"
        "  └─ _Visualizar:_ `/reunioes <id>` para ver os detalhes completos.\n"
        "• `/transcrever` - Lista arquivos de áudio em `uploads/` para transcrição.\n"
        "  └─ _Transcrever:_ `/transcrever <nome_do_arquivo>` para iniciar.\n"
        "• `/backup` - Cria uma cópia de segurança criptografada do banco.\n"
        "• `/restore` - Restaura um backup criptografado.\n"
    )

def format_notes_telegram() -> str:
    notes = db.list_all_user_notes()
    if not notes:
        return "📝 *Memória de Longo Prazo*\n\nNenhuma anotação salva."
        
    linhas = ["📝 *Suas Notas Salvas (Memória de Longo Prazo):*\n"]
    for nid, content, dt in notes:
        dt_str = dt.strftime("%d/%m/%Y %H:%M")
        display_content = content[:200] + "..." if len(content) > 200 else content
        linhas.append(f"• *#{nid}* ({dt_str}):\n  _{display_content}_")
        
    linhas.append("\n_Use o chat normal para fazer perguntas sobre suas notas ou pesquisar por elas._")
    return "\n".join(linhas)

def format_cron_telegram() -> str:
    jobs = db.get_active_cron_jobs()
    if not jobs:
        return "🤖 *Subagentes Agendados (Cron Jobs)*\n\nNenhuma tarefa agendada ativa."
        
    linhas = ["🤖 *Subagentes Agendados (Cron Jobs):*\n"]
    for j in jobs:
        last_str = j["last_run"].strftime("%d/%m/%Y %H:%M") if j["last_run"] else "Nunca"
        next_str = j["next_run"].strftime("%d/%m/%Y %H:%M") if j["next_run"] else "N/A"
        p_desc = j["task_prompt"][:100] + "..." if len(j["task_prompt"]) > 100 else j["task_prompt"]
        
        linhas.append(
            f"• *#{j['id']} - {j['name']}*\n"
            f"  📅 *Cron:* `{j['cron_expression']}`\n"
            f"  🔔 *Próximo Disparo:* `{next_str}` | *Último:* `{last_str}`\n"
            f"  💬 *Prompt:* _{p_desc}_"
        )
        
    return "\n\n".join(linhas)

def format_finance_card_list() -> str:
    cards = db.get_credit_cards().get("cartoes", {})
    if not cards:
        return "💳 *Cartões de Crédito*\n\nNenhum cartão cadastrado. Use `/finance card add` no console ou no bot."
        
    linhas = ["💳 *Cartões de Crédito Cadastrados:*\n"]
    for name, info in cards.items():
        linhas.append(f"• *{name}* - Fechamento: dia `{info['closing_day']}` | Vencimento: dia `{info['due_day']}`")
        
    return "\n".join(linhas)

def format_finance_telegram(parts: list) -> str:
    from datetime import datetime
    now = datetime.now()
    current_month_str = now.strftime("%m-%Y")
    
    if now.month == 12:
        next_month = 1
        next_year = now.year + 1
    else:
        next_month = now.month + 1
        next_year = now.year
    next_month_str = f"{next_month:02d}-{next_year}"
    
    limit = None
    month_year = None
    query = None
    table_title = ""
    show_deleted = False
    
    if len(parts) == 1:
        month_year = current_month_str
        table_title = f"Mês Atual ({current_month_str})"
    else:
        arg = parts[1]
        if arg.lower() == "all":
            table_title = "Todas as Transações"
        elif arg.lower() == "next":
            month_year = next_month_str
            table_title = f"Próximo Mês ({next_month_str})"
        elif arg.lower() == "deleted":
            show_deleted = True
            table_title = "Transações Inativas"
        elif arg.lower().startswith("mes="):
            month_year = arg.split("=")[1]
            table_title = f"Vencimento em {month_year}"
        elif arg.lower().startswith("q="):
            query = arg.split("=")[1]
            table_title = f"Busca por '{query}'"
        else:
            return (
                "❌ *Filtro inválido.*\nUse:\n"
                "• `/finance` (mês atual)\n"
                "• `/finance next` (próximo mês)\n"
                "• `/finance all` (tudo)\n"
                "• `/finance deleted` (excluídos)\n"
                "• `/finance mes=MM-YYYY` (mês específico)\n"
                "• `/finance q=busca` (pesquisar)"
            )
            
    if show_deleted:
        records = db.get_deleted_financial_records()
    else:
        records = db.search_financial_records(limit=limit, month_year=month_year, query=query)
        
    if not records:
        return f"🔍 *{table_title}*\n\nNenhuma transação encontrada com os filtros especificados."
        
    sum_receitas = 0.0
    sum_despesas = 0.0
    for _, rtype, _, val, _, _, _ in records:
        if rtype == "receita":
            sum_receitas += val
        elif rtype == "despesa":
            sum_despesas += val
    sum_saldo = sum_receitas - sum_despesas
    
    resumo = (
        f"📊 *Resumo Financeiro - {table_title}*\n"
        f"🟢 Receitas: `R$ {sum_receitas:.2f}`\n"
        f"🔴 Despesas: `R$ {sum_despesas:.2f}`\n"
        f"⚖️ Saldo: `R$ {sum_saldo:.2f}`\n"
        f"─────────────────────\n"
    )
    
    linhas = []
    for rid, rtype, cat, val, desc, dt, due_dt in records:
        emoji = "🟢" if rtype == "receita" else "🔴"
        dt_str = dt.strftime("%d/%m/%Y")
        due_str = due_dt.strftime("%d/%m/%Y") if due_dt else "N/A"
        desc_str = f" - _{desc}_" if desc else ""
        
        linha = (
            f"• *#{rid}* ({dt_str}) | {emoji} {cat}\n"
            f"  *Valor:* `R$ {val:.2f}` | *Venc:* `{due_str}`{desc_str}"
        )
        linhas.append(linha)
        
    return resumo + "\n\n".join(linhas)

def format_markdown_for_telegram(text: str) -> str:
    """
    Converte markdown padrão da LLM para o formato aceito pelo Telegram Markdown (v1).
    - Substitui '**' por '*' para negrito, preservando blocos de código.
    """
    if not text:
        return text
        
    code_blocks = []
    def replace_code_block(match):
        code_blocks.append(match.group(0))
        return f"__CODE_BLOCK_{len(code_blocks)-1}__"
        
    # Regex para capturar blocos de código ```...``` e `...`
    text_temp = re.sub(r'```[\s\S]*?```|`.*?`', replace_code_block, text)
    
    # Converte negrito de '**' para '*'
    text_temp = text_temp.replace('**', '*')
    
    # Restaura blocos de código
    for i, block in enumerate(code_blocks):
        text_temp = text_temp.replace(f"__CODE_BLOCK_{i}__", block)
        
    return text_temp

def reply_formatted(message, text):
    """Envia a resposta formatada para o Telegram, caindo para texto puro se o parser falhar."""
    try:
        formatted_text = format_markdown_for_telegram(text)
        bot.reply_to(message, formatted_text, parse_mode="Markdown")
    except Exception as e:
        logging.warning("Falha ao renderizar Markdown no Telegram: %s", e)
        try:
            bot.reply_to(message, text)
        except Exception as e_fallback:
            logging.error("Falha crítica ao enviar mensagem: %s", e_fallback)

@bot.message_handler(func=lambda msg: True)
def handle_incoming_message(message):
    sender_username = message.from_user.username or "SemUsername"
    logging.info("Mensagem recebida de @%s (ID: %s): '%s'", sender_username, message.chat.id, message.text)
    
    if not is_authorized(message):
        logging.warning("Tentativa de acesso não autorizada de Chat ID: %s (@%s)", message.chat.id, sender_username)
        bot.reply_to(message, "Acesso negado. Este bot do agente é privado.")
        return
        
    text = message.text.strip()
    if not text:
        return
        
    # Verifica se o usuário estava aguardando enviar instrução personalizada para áudio
    if user_states.get(message.chat.id) == "waiting_for_custom_instruction":
        process_custom_audio_instruction(message, text)
        return
        
    # Tratamento de Comandos Especiais com Barra
    if text.startswith("/"):
        if text in ["/exit", "/quit"]:
            bot.reply_to(message, "Comandos de finalização (/exit, /quit) são desativados via chat remoto.")
            return
            
        parts = text.split()
        cmd = parts[0].lower()
        
        # Comandos permitidos sem login ativo
        if cmd not in ("/login", "/help", "/start"):
            if not is_logged_in(message):
                return
                
        bot.send_chat_action(message.chat.id, 'typing')
        
        # Filtros e roteamentos personalizados para exibir layouts bonitos no celular
        try:
            if cmd == "/help" or cmd == "/start":
                bot.reply_to(message, format_help_telegram(), parse_mode="Markdown")
                return
            elif cmd == "/notes":
                bot.reply_to(message, format_notes_telegram(), parse_mode="Markdown")
                return
            elif cmd == "/cron" and len(parts) == 1:
                bot.reply_to(message, format_cron_telegram(), parse_mode="Markdown")
                return
            elif cmd == "/reunioes":
                handle_reunioes_command(message, parts)
                return
            elif cmd == "/transcrever":
                handle_transcrever_command(message, parts)
                return
            elif cmd == "/models":
                handle_models_command(message, parts)
                return
            elif cmd == "/finance":
                # Verifica se é uma listagem/filtro normal ou uma ação (delete, restore, import, card add/buy)
                is_list = True
                if len(parts) > 1:
                    sub = parts[1].lower()
                    if sub in ["delete", "restore", "import", "csv"] or (sub in ["card", "cartao"] and len(parts) > 2 and parts[2].lower() in ["add", "buy"]):
                        is_list = False
                        
                if is_list:
                    # Verifica se é listagem de cartões de crédito
                    if len(parts) > 1 and parts[1].lower() in ["card", "cartao"] and (len(parts) == 2 or parts[2].lower() == "list"):
                        bot.reply_to(message, format_finance_card_list(), parse_mode="Markdown")
                    else:
                        bot.reply_to(message, format_finance_telegram(parts), parse_mode="Markdown")
                    return
            
            # Caso contrário, cai no fallback de capturar o console Rich do terminal
            from meu_agente_cli.main import console
            
            with console.capture() as capture:
                handle_slash_command(text)
                
            output = capture.get()
            clean_output = strip_ansi(output).strip()
            
            if not clean_output:
                clean_output = "Comando executado sem retorno textual."
                
            # Formata tabelas ou desenhos unicode em monospace no Telegram
            if any(char in clean_output for char in ["┏", "┃", "─", "┼", "│", "┌", "╭", "╰"]):
                if len(clean_output) > 4000:
                    clean_output = clean_output[:3900] + "\n... (saída longa truncada)"
                bot.reply_to(message, f"```\n{clean_output}\n```", parse_mode="Markdown")
            else:
                if len(clean_output) > 4000:
                    clean_output = clean_output[:3900] + "\n... (saída longa truncada)"
                bot.reply_to(message, clean_output)
        except Exception as e:
            logging.exception("Erro ao executar comando: %s", text)
            bot.reply_to(message, f"Erro ao executar comando: {e}")
            
    else:
        # Conversa normal com o Agente (Pensamento + Execução Silenciosa de Ferramentas)
        if not is_logged_in(message):
            return
            
        bot.send_chat_action(message.chat.id, 'typing')
        try:
            telegram_prompt = f"{text}{TELEGRAM_INSTRUCTION}"
            response = agent.process_agent_turn_silent(telegram_prompt)
            reply_formatted(message, response)
        except Exception as e:
            logging.exception("Erro no processamento da conversa para a mensagem: %s", text)
            bot.reply_to(message, f"Erro no processamento da conversa: {e}")

def escape_markdown(text: str) -> str:
    """Escapa caracteres especiais do Markdown V1 do Telegram para evitar quebras de parser."""
    if not text:
        return text
    for char in ['_', '*', '`']:
        text = text.replace(char, f"\\{char}")
    return text

# Helper functions for Audio processing & History
def process_custom_audio_instruction(message, instruction):
    chat_id = message.chat.id
    user_states.pop(chat_id, None)  # Limpa o estado
    
    trans_data = pending_transcriptions.get(chat_id)
    if not trans_data:
        bot.reply_to(message, "⚠️ Nenhuma transcrição ativa encontrada. Envie o áudio novamente.")
        return
        
    transcription = trans_data["text"]
    audio_path = trans_data["audio_path"]
    
    bot.send_chat_action(chat_id, 'typing')
    bot.reply_to(message, "✍️ Processando transcrição com sua instrução personalizada...")
    
    try:
        prompt = (
            f"Processar a seguinte transcrição de áudio seguindo a instrução: {instruction}\n\n"
            f"Transcrição:\n{transcription}"
            f"{TELEGRAM_INSTRUCTION}"
        )
        response = agent.process_agent_turn_silent(prompt)
        
        # Salva permanentemente no banco
        db.save_audio_transcription(chat_id, audio_path, transcription, f"Instrução personalizada: {instruction}", response)
        
        reply_formatted(message, response)
        # Limpa transcrição temporária
        pending_transcriptions.pop(chat_id, None)
    except Exception as e:
        bot.reply_to(message, f"❌ Erro ao processar áudio: {e}")

def handle_transcrever_command(message, parts):
    chat_id = message.chat.id
    uploads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads"))
    os.makedirs(uploads_dir, exist_ok=True)
    
    allowed_exts = {".mp3", ".wav", ".m4a", ".ogg", ".oga", ".aac", ".flac", ".wma"}
    
    # Lista arquivos
    files = []
    if os.path.exists(uploads_dir):
        files = [f for f in os.listdir(uploads_dir) if os.path.isfile(os.path.join(uploads_dir, f)) and os.path.splitext(f)[1].lower() in allowed_exts]
        
    if len(parts) == 1:
        if not files:
            bot.reply_to(
                message,
                "📂 *Pasta de uploads vazia*\n\n"
                "Nenhum arquivo de áudio foi encontrado no diretório `uploads/`.\n"
                "Coloque arquivos de áudio manualmente no servidor e digite `/transcrever <nome_do_arquivo>`.",
                parse_mode="Markdown"
            )
            return
            
        linhas = ["📂 *Arquivos de áudio disponíveis para transcrição:*"]
        for f in files:
            linhas.append(f"• `{f}`")
        linhas.append("\n💡 Digite `/transcrever <nome_do_arquivo>` para iniciar o processamento.")
        bot.reply_to(message, "\n".join(linhas), parse_mode="Markdown")
        return
        
    # Processamento do arquivo especificado
    filename = " ".join(parts[1:])  # Para o caso do nome ter espaços
    file_path = os.path.join(uploads_dir, filename)
    
    if not os.path.exists(file_path):
        bot.reply_to(message, f"❌ O arquivo `{filename}` não foi encontrado em `uploads/`.")
        return
        
    bot.send_chat_action(chat_id, 'record_audio')
    bot.reply_to(message, f"📥 Iniciando o processamento do arquivo `{filename}`... Isso pode levar alguns minutos dependendo do tamanho.")
    
    try:
        # Cria diretório de histórico
        archive_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads", "archive", "audio"))
        os.makedirs(archive_dir, exist_ok=True)
        
        timestamp = int(time.time())
        ext = os.path.splitext(filename)[1]
        
        original_filename = f"{timestamp}_{filename}"
        original_filepath = os.path.join(archive_dir, original_filename)
        
        wav_filename = f"{timestamp}_{os.path.splitext(filename)[0]}.wav"
        wav_filepath = os.path.join(archive_dir, wav_filename)
        
        # Copia o arquivo original para a pasta de histórico
        import shutil
        shutil.copy(file_path, original_filepath)
        
        # Converte para WAV com ffmpeg
        cmd = ["ffmpeg", "-y", "-i", original_filepath, "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000", wav_filepath]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        if result.returncode != 0:
            error_msg = result.stderr.decode('utf-8', errors='ignore')
            print(f"[ERROR] Falha na conversão de áudio pelo ffmpeg: {error_msg}")
            bot.reply_to(message, f"❌ Erro ao converter o arquivo `{filename}` para WAV. Verifique o formato.")
            return
            
        # Faz a transcrição
        from meu_agente_cli.speech_to_text import transcrever_arquivo_audio
        
        bot.send_chat_action(chat_id, 'typing')
        transcription = transcrever_arquivo_audio(wav_filepath)
        
        if not transcription.strip() or transcription.startswith("["):
            bot.reply_to(message, f"⚠️ Não consegui extrair palavras compreensíveis do arquivo `{filename}`.\n\nRetorno: {transcription or 'Áudio em silêncio'}")
            return
            
        # Salva no estado temporário
        pending_transcriptions[chat_id] = {
            "text": transcription,
            "audio_path": wav_filepath
        }
        
        # Monta teclado inline
        markup = InlineKeyboardMarkup()
        btn_simple = InlineKeyboardButton("📝 Resumo Simples", callback_data="audio_action:simple")
        btn_detailed = InlineKeyboardButton("📌 Pontos & Plano de Ação", callback_data="audio_action:detailed")
        btn_custom = InlineKeyboardButton("✍️ Outra Instrução...", callback_data="audio_action:custom")
        
        markup.row(btn_simple, btn_detailed)
        markup.row(btn_custom)
        
        # Resposta
        transcription_escaped = escape_markdown(transcription[:1500])
        msg_text = (
            f"🎙️ *Áudio Transcrito com Sucesso (`{filename}`):*\n\n"
            f"_\"{transcription_escaped}... (exibindo primeiros 1500 caracteres)\"_\n\n"
            f"Escolha o que deseja fazer com esta transcrição:"
        )
        try:
            bot.reply_to(message, msg_text, reply_markup=markup, parse_mode="Markdown")
        except Exception as e:
            print(f"[Warning] Falha ao renderizar Markdown em /transcrever response: {e}")
            try:
                msg_text_plain = (
                    f"🎙️ Áudio Transcrito com Sucesso ({filename}):\n\n"
                    f"\"{transcription[:1500]}... (exibindo primeiros 1500 caracteres)\"\n\n"
                    f"Escolha o que deseja fazer com esta transcrição:"
                )
                bot.reply_to(message, msg_text_plain, reply_markup=markup)
            except Exception as e_fallback:
                print(f"[ERROR] Falha crítica ao enviar resposta de transcrição: {e_fallback}")
        
    except Exception as e:
        print(f"[ERROR] Falha geral ao transcrever arquivo manual: {e}")
        bot.reply_to(message, f"❌ Erro no processamento do arquivo de áudio: {e}")

def handle_models_command(message, parts):
    chat_id = message.chat.id
    
    if len(parts) == 1:
        llm_provider = db.get_setting("llm_provider", "lm_studio")
        active_model = db.get_setting("active_model", "Nenhum")
        
        linhas = [
            "🤖 *Configuração do Modelo de Linguagem (LLM)*",
            f"• Provedor atual: *{llm_provider.upper()}*",
            f"• Modelo atual: *{active_model}*\n",
            "💡 *Como alterar o modelo via Telegram:*",
            "1️⃣ *LM Studio (Local)*:",
            "   └─ Digite `/models lm_studio` para ver os modelos carregados localmente.",
            "   └─ Digite `/models lm_studio <nome_do_modelo>` para ativar.",
            "",
            "2️⃣ *Provedores Externos*:",
            "   └─ Digite `/models <provedor> <API_KEY> [modelo]`",
            "   └─ *Provedores suportados:* `openai`, `gemini`, `claude`, `deepseek`, `qwen`, `kimi`",
            "   └─ *Exemplo:* `/models openai sk-proj-... gpt-4o-mini`",
            "",
            "3️⃣ *Provedor Personalizado (OpenAI-Compatible)*:",
            "   └─ Digite `/models custom <API_KEY> <URL_BASE> <modelo>`",
            "   └─ *Exemplo:* `/models custom sk-key http://192.168.1.50:8000/v1 meu-modelo-lhamas`"
        ]
        bot.reply_to(message, "\n".join(linhas), parse_mode="Markdown")
        return
        
    provider = parts[1].lower()
    
    if provider == "lm_studio":
        from meu_agente_cli import llm
        models = llm.get_available_models()
        if len(parts) == 2:
            if not models:
                bot.reply_to(message, "❌ Nenhum modelo ativo detectado no LM Studio local. Certifique-se de que o LM Studio está aberto e com um modelo carregado.")
                return
                
            linhas = ["🤖 *Modelos disponíveis no LM Studio:*"]
            for m in models:
                linhas.append(f"• `{m}`")
            linhas.append("\n👉 Para ativar um deles, digite: `/models lm_studio <nome_do_modelo>`")
            bot.reply_to(message, "\n".join(linhas), parse_mode="Markdown")
            return
        else:
            model_name = " ".join(parts[2:])
            if model_name not in models:
                # Se não bater exatamente, verifica se é substring ou aceita direto
                matched = [m for m in models if model_name.lower() in m.lower()]
                if len(matched) == 1:
                    model_name = matched[0]
                elif len(matched) > 1:
                    bot.reply_to(message, f"❓ Modelo ambíguo. Encontrados:\n" + "\n".join([f"• `{m}`" for m in matched]))
                    return
            
            db.set_setting("llm_provider", "lm_studio")
            db.set_setting("active_model", model_name)
            bot.reply_to(message, f"✅ Provedor alterado para *LM Studio* e modelo ativo para `{model_name}`.", parse_mode="Markdown")
            return
            
    elif provider in ["openai", "gemini", "claude", "deepseek", "qwen", "kimi"]:
        if len(parts) < 3:
            bot.reply_to(message, f"❌ Uso correto: `/models {provider} <API_KEY> [modelo]`")
            return
            
        api_key = parts[2]
        
        default_models = {
            "openai": "gpt-4o-mini",
            "gemini": "gemini-1.5-flash",
            "claude": "claude-3-5-sonnet-latest",
            "deepseek": "deepseek-chat",
            "qwen": "qwen-plus",
            "kimi": "moonshot-v1-8k"
        }
        
        model_name = parts[3] if len(parts) > 3 else default_models.get(provider, "")
        
        db.set_setting("llm_provider", provider)
        db.set_setting("provider_api_key", api_key)
        db.set_setting("active_model", model_name)
        
        bot.reply_to(message, f"✅ Provedor alterado para *{provider.upper()}*.\n🔑 API Key salva no banco de dados.\n🤖 Modelo ativo: `{model_name}`.", parse_mode="Markdown")
        return
        
    elif provider == "custom":
        if len(parts) < 5:
            bot.reply_to(message, "❌ Uso correto: `/models custom <API_KEY> <URL_BASE> <modelo>`\nUse `none` na API KEY se não for necessária autenticação.")
            return
            
        api_key = parts[2]
        base_url = parts[3]
        model_name = " ".join(parts[4:])
        
        if api_key.lower() == "none":
            api_key = ""
            
        db.set_setting("llm_provider", "custom")
        db.set_setting("provider_api_key", api_key)
        db.set_setting("provider_base_url", base_url)
        db.set_setting("active_model", model_name)
        
        bot.reply_to(message, f"✅ Provedor alterado para *CUSTOM*.\n🌐 URL Base: `{base_url}`\n🤖 Modelo ativo: `{model_name}`.", parse_mode="Markdown")
        return
        
    else:
        bot.reply_to(message, "❌ Provedor desconhecido. Use `/models` para ver a lista de opções.")

def handle_reunioes_command(message, parts):
    chat_id = message.chat.id
    if len(parts) == 1:
        # Listagem ativa
        rows = db.list_audio_transcriptions(chat_id)
        if not rows:
            bot.reply_to(message, "📂 *Histórico de Reuniões*\n\nNenhuma transcrição de áudio/reunião ativa encontrada no banco.")
            return
            
        linhas = ["📂 *Gravações e Reuniões Salvas:*"]
        for r_id, audio_path, transcription, user_prompt, created_at in rows:
            dt_str = created_at.strftime("%d/%m/%Y %H:%M")
            trecho = transcription[:65] + "..." if len(transcription) > 65 else transcription
            linhas.append(f"• *#{r_id}* ({dt_str}) | Foco: _{user_prompt}_\n  _\"{trecho}\"_")
            
        linhas.append("\n_Use `/reunioes <id>` para ver os detalhes, `/reunioes delete <id>` para inativar, ou `/reunioes deleted` para ver inativas._")
        bot.reply_to(message, "\n".join(linhas), parse_mode="Markdown")
        
    elif len(parts) >= 2:
        subcmd = parts[1].lower()
        
        if subcmd == "deleted":
            # Listagem de deletadas
            rows = db.list_deleted_audio_transcriptions(chat_id)
            if not rows:
                bot.reply_to(message, "📂 *Histórico de Reuniões Inativas*\n\nNenhuma reunião inativada encontrada.")
                return
                
            linhas = ["📂 *Reuniões Inativadas (Soft Deleted):*"]
            for r_id, audio_path, transcription, user_prompt, created_at in rows:
                dt_str = created_at.strftime("%d/%m/%Y %H:%M")
                trecho = transcription[:65] + "..." if len(transcription) > 65 else transcription
                linhas.append(f"• *#{r_id}* ({dt_str}) | Foco: _{user_prompt}_\n  _\"{trecho}\"_")
                
            linhas.append("\n_Use `/reunioes restore <id>` para reativar uma reunião._")
            bot.reply_to(message, "\n".join(linhas), parse_mode="Markdown")
            return
            
        elif subcmd == "delete":
            if len(parts) < 3:
                bot.reply_to(message, "❌ Por favor, especifique o ID. Exemplo: `/reunioes delete 5`.")
                return
            try:
                doc_id = int(parts[2])
            except ValueError:
                bot.reply_to(message, "❌ ID inválido. Use `/reunioes delete <id>` com o número do registro.")
                return
                
            # Verifica se existe
            row = db.get_audio_transcription_by_id(chat_id, doc_id)
            if not row:
                bot.reply_to(message, f"❌ Reunião com ID *#{doc_id}* não encontrada.")
                return
                
            if db.delete_audio_transcription(chat_id, doc_id):
                bot.reply_to(message, f"🗑️ Reunião *#{doc_id}* foi inativada com sucesso (Soft Deleted).\n\nVocê pode recuperá-la digitando `/reunioes restore {doc_id}`.")
            else:
                bot.reply_to(message, f"❌ Falha ao inativar reunião *#{doc_id}*.")
            return
            
        elif subcmd == "restore":
            if len(parts) < 3:
                bot.reply_to(message, "❌ Por favor, especifique o ID. Exemplo: `/reunioes restore 5`.")
                return
            try:
                doc_id = int(parts[2])
            except ValueError:
                bot.reply_to(message, "❌ ID inválido. Use `/reunioes restore <id>` com o número do registro.")
                return
                
            # Verifica se existe
            row = db.get_audio_transcription_by_id(chat_id, doc_id)
            if not row:
                bot.reply_to(message, f"❌ Reunião com ID *#{doc_id}* não encontrada.")
                return
                
            if db.restore_audio_transcription(chat_id, doc_id):
                bot.reply_to(message, f"✅ Reunião *#{doc_id}* foi reativada com sucesso!")
            else:
                bot.reply_to(message, f"❌ Falha ao reativar reunião *#{doc_id}*.")
            return
            
        else:
            # Caso seja `/reunioes <id>`
            try:
                doc_id = int(parts[1])
            except ValueError:
                bot.reply_to(message, "❌ Subcomando ou ID inválido. Use `/reunioes`, `/reunioes <id>`, `/reunioes delete <id>`, ou `/reunioes restore <id>`.")
                return
                
            row = db.get_audio_transcription_by_id(chat_id, doc_id)
            if not row:
                bot.reply_to(message, f"❌ Registro de reunião com ID *#{doc_id}* não encontrado.")
                return
                
            r_id, audio_path, transcription, user_prompt, llm_response, created_at, active = row
            dt_str = created_at.strftime("%d/%m/%Y %H:%M")
            status_text = "🟢 Ativa" if active else "🔴 Inativa (Soft Deleted)"
            
            # Escapa campos de entrada do usuário para evitar quebra de Markdown
            user_prompt_escaped = escape_markdown(user_prompt)
            transcription_escaped = escape_markdown(transcription)
            
            texto_completo = (
                f"📅 *Detalhes da Reunião #{r_id}* ({dt_str})\n"
                f"📊 *Status:* {status_text}\n"
                f"🔊 *Áudio:* `{os.path.basename(audio_path)}`\n"
                f"🎯 *Solicitação:* _{user_prompt_escaped}_\n\n"
                f"─────────────────────\n"
                f"🎙️ *Transcrição Completa:*\n"
                f"_{transcription_escaped}_\n\n"
                f"─────────────────────\n"
                f"🤖 *Resultado / Resumo:*\n\n"
                f"{llm_response}"
            )
            
            # Envia de forma estruturada, tratando possíveis erros de renderização de Markdown
            try:
                if len(texto_completo) > 4000:
                    parts_txt = [texto_completo[i:i+3900] for i in range(0, len(texto_completo), 3900)]
                    for pt in parts_txt:
                        bot.send_message(chat_id, pt, parse_mode="Markdown")
                else:
                    bot.reply_to(message, texto_completo, parse_mode="Markdown")
            except Exception as e:
                print(f"[Warning] Falha ao renderizar Markdown em /reunioes {doc_id}: {e}")
                # Fallback: envia o texto puro sem parser
                try:
                    if len(texto_completo) > 4000:
                        parts_txt = [texto_completo[i:i+3900] for i in range(0, len(texto_completo), 3900)]
                        for pt in parts_txt:
                            bot.send_message(chat_id, pt)
                    else:
                        bot.reply_to(message, texto_completo)
                except Exception as e_fallback:
                    print(f"[ERROR] Falha crítica ao enviar detalhes da reunião #{doc_id}: {e_fallback}")

@bot.message_handler(content_types=['voice', 'audio'])
def handle_audio_upload(message):
    if not is_authorized(message):
        print(f"[BLOQUEADO] Áudio recebido de Chat ID não autorizado: {message.chat.id}")
        bot.reply_to(message, "Acesso negado. Este bot do agente é privado.")
        return

    if not is_logged_in(message):
        return

    bot.send_chat_action(message.chat.id, 'record_audio')
    
    try:
        # Identifica se é voice (mensagem de voz gravada) ou audio (arquivo anexado)
        is_voice = message.content_type == 'voice'
        media = message.voice if is_voice else message.audio
        
        # Validação do limite de tamanho do Telegram para download de arquivos por bots (20 MB)
        if media.file_size and media.file_size > 20 * 1024 * 1024:
            bot.reply_to(
                message,
                "⚠️ *Arquivo de áudio muito grande!*\n\n"
                "O Telegram limita o download de arquivos por bots a no máximo *20 MB*.\n"
                "Seu arquivo possui aproximadamente *{:.1f} MB*.\n\n"
                "💡 *Sugestão:* Divida o áudio em partes menores de até 20 MB ou converta-o para um formato mais compactado."
                .format(media.file_size / (1024 * 1024)),
                parse_mode="Markdown"
            )
            return
            
        file_id = media.file_id
        
        # Obtém caminhos
        file_info = bot.get_file(file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        # Caminho do diretório uploads/archive/audio/
        archive_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads", "archive", "audio"))
        os.makedirs(archive_dir, exist_ok=True)
        
        # Define nomes de arquivos
        timestamp = int(time.time())
        ext = os.path.splitext(file_info.file_path)[1]
        if not ext:
            ext = ".ogg" if is_voice else ".mp3"
            
        original_filename = f"{timestamp}_{file_id}{ext}"
        original_filepath = os.path.join(archive_dir, original_filename)
        
        wav_filename = f"{timestamp}_{file_id}.wav"
        wav_filepath = os.path.join(archive_dir, wav_filename)
        
        # Salva o arquivo de áudio original recebido
        with open(original_filepath, 'wb') as f:
            f.write(downloaded_file)
        
        bot.reply_to(message, "📥 Áudio recebido! Convertendo e preparando para transcrição...")
        
        # Conversão via ffmpeg para WAV
        cmd = ["ffmpeg", "-y", "-i", original_filepath, "-acodec", "pcm_s16le", "-ac", "1", "-ar", "16000", wav_filepath]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        if result.returncode != 0:
            error_msg = result.stderr.decode('utf-8', errors='ignore')
            print(f"[ERROR] Falha na conversão de áudio pelo ffmpeg: {error_msg}")
            bot.reply_to(message, "❌ Erro ao converter o arquivo de áudio. Verifique se o formato está correto.")
            return
            
        # Transcrição
        from meu_agente_cli.speech_to_text import transcrever_arquivo_audio
        
        bot.send_chat_action(message.chat.id, 'typing')
        transcription = transcrever_arquivo_audio(wav_filepath)
        
        if not transcription.strip() or transcription.startswith("["):
            bot.reply_to(message, f"⚠️ Não consegui extrair palavras compreensíveis do áudio.\n\nRetorno: {transcription or 'Áudio em silêncio'}")
            return
            
        # Guarda no estado temporário (em memória)
        pending_transcriptions[message.chat.id] = {
            "text": transcription,
            "audio_path": wav_filepath
        }
        
        # Monta teclado inline com opções
        markup = InlineKeyboardMarkup()
        btn_simple = InlineKeyboardButton("📝 Resumo Simples", callback_data="audio_action:simple")
        btn_detailed = InlineKeyboardButton("📌 Pontos & Plano de Ação", callback_data="audio_action:detailed")
        btn_custom = InlineKeyboardButton("✍️ Outra Instrução...", callback_data="audio_action:custom")
        
        markup.row(btn_simple, btn_detailed)
        markup.row(btn_custom)
        
        # Envia resposta
        transcription_escaped = escape_markdown(transcription)
        msg_text = (
            f"🎙️ *Áudio Transcrito com Sucesso:*\n\n"
            f"_\"{transcription_escaped}\"_\n\n"
            f"Escolha o que deseja fazer com esta transcrição:"
        )
        try:
            bot.reply_to(message, msg_text, reply_markup=markup, parse_mode="Markdown")
        except Exception as e:
            print(f"[Warning] Falha ao renderizar Markdown em handle_audio_upload response: {e}")
            try:
                msg_text_plain = (
                    f"🎙️ Áudio Transcrito com Sucesso:\n\n"
                    f"\"{transcription}\"\n\n"
                    f"Escolha o que deseja fazer com esta transcrição:"
                )
                bot.reply_to(message, msg_text_plain, reply_markup=markup)
            except Exception as e_fallback:
                print(f"[ERROR] Falha crítica ao enviar resposta de transcrição de voz: {e_fallback}")
        
    except Exception as e:
        print(f"[ERROR] Falha geral no recebimento e transcrição de áudio: {e}")
        bot.reply_to(message, f"❌ Ocorreu um erro no processamento do áudio: {e}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("audio_action:"))
def handle_audio_action(call):
    chat_id = call.message.chat.id
    action = call.data.split(":")[1]
    
    trans_data = pending_transcriptions.get(chat_id)
    if not trans_data:
        bot.answer_callback_query(call.id, "Nenhuma transcrição ativa encontrada.")
        bot.send_message(chat_id, "⚠️ Desculpe, a transcrição expirou ou não foi encontrada. Envie o áudio novamente.")
        return
        
    transcription = trans_data["text"]
    audio_path = trans_data["audio_path"]
    
    bot.answer_callback_query(call.id)
    
    if action == "simple":
        bot.send_message(chat_id, "📝 Gerando resumo simples...")
        bot.send_chat_action(chat_id, 'typing')
        try:
            prompt = f"Faça um resumo simples e direto da seguinte transcrição de áudio:\\n\\n{transcription}{TELEGRAM_INSTRUCTION}"
            response = agent.process_agent_turn_silent(prompt)
            
            db.save_audio_transcription(chat_id, audio_path, transcription, "Resumo Simples", response)
            
            bot.send_message(chat_id, response, parse_mode="Markdown")
            pending_transcriptions.pop(chat_id, None)
        except Exception as e:
            bot.send_message(chat_id, f"❌ Erro ao gerar resumo: {e}")
            
    elif action == "detailed":
        bot.send_message(chat_id, "📌 Gerando principais pontos e plano de ação...")
        bot.send_chat_action(chat_id, 'typing')
        try:
            prompt = (
                f"Analise a seguinte transcrição de áudio e forneça um resumo estruturado contendo "
                f"os principais pontos discutidos e um plano de ação detalhado com os próximos passos:\n\n"
                f"{transcription}{TELEGRAM_INSTRUCTION}"
            )
            response = agent.process_agent_turn_silent(prompt)
            
            db.save_audio_transcription(chat_id, audio_path, transcription, "Principais Pontos e Plano de Ação", response)
            
            bot.send_message(chat_id, response, parse_mode="Markdown")
            pending_transcriptions.pop(chat_id, None)
        except Exception as e:
            bot.send_message(chat_id, f"❌ Erro ao gerar resumo: {e}")
            
    elif action == "custom":
        user_states[chat_id] = "waiting_for_custom_instruction"
        bot.send_message(chat_id, "✍️ Digite a instrução personalizada para processar o áudio (ex: 'traduza para inglês' ou 'extraia datas'):")

def extract_pdf_text(file_path: str) -> str:
    """Extrai texto de um arquivo PDF usando a biblioteca pypdf."""
    import pypdf
    texto = []
    try:
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                texto.append(f"--- Página {i+1} ---\n{page_text}")
        return "\n".join(texto).strip()
    except Exception as e:
        return f"[Erro ao ler PDF: {str(e)}]"

def extract_docx_text(file_path: str) -> str:
    """Extrai texto de um arquivo do Word (.docx) usando python-docx."""
    import docx
    try:
        doc = docx.Document(file_path)
        texto = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                texto.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                texto.append(" | ".join(row_text))
        return "\n".join(texto).strip()
    except Exception as e:
        return f"[Erro ao ler arquivo Word: {str(e)}]"

def extract_xlsx_text(file_path: str) -> str:
    """Extrai dados de planilhas Excel (.xlsx) formatando como tabelas legíveis."""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        sheets_data = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheets_data.append(f"=== Planilha: {sheet_name} ===")
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    sheets_data.append(row_str)
        return "\n".join(sheets_data).strip()
    except Exception as e:
        return f"[Erro ao ler planilha Excel: {str(e)}]"

@bot.message_handler(content_types=['document'])
def handle_document_upload(message):
    if not is_authorized(message):
        print(f"[BLOQUEADO] Documento recebido de Chat ID não autorizado: {message.chat.id}")
        bot.reply_to(message, "Acesso negado. Este bot do agente é privado.")
        return

    if not is_logged_in(message):
        return

    bot.send_chat_action(message.chat.id, 'typing')
    
    try:
        # Validação do limite de tamanho do Telegram para download de arquivos por bots (20 MB)
        if message.document.file_size and message.document.file_size > 20 * 1024 * 1024:
            bot.reply_to(
                message,
                "⚠️ *Arquivo muito grande!*\n\n"
                "O Telegram limita o download de arquivos por bots a no máximo *20 MB*.\n"
                "Seu arquivo possui aproximadamente *{:.1f} MB*.\n\n"
                "💡 *Sugestão:* Por favor, envie um arquivo menor que 20 MB."
                .format(message.document.file_size / (1024 * 1024)),
                parse_mode="Markdown"
            )
            return
            
        file_info = bot.get_file(message.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        # Cria a pasta uploads/ se não existir
        uploads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads"))
        os.makedirs(uploads_dir, exist_ok=True)
        
        file_name = message.document.file_name
        file_path = os.path.join(uploads_dir, file_name)
        
        with open(file_path, 'wb') as new_file:
            new_file.write(downloaded_file)
            
        print(f"[DOCUMENTO] Salvo com sucesso em: {file_path}")
        
        # Se for um CSV para importação financeira
        if file_name.lower().endswith(".csv"):
            bot.reply_to(message, f"📥 Arquivo CSV `{file_name}` recebido. Iniciando a importação financeira...")
            dest_csv = os.path.join(uploads_dir, "finance.csv")
            import shutil
            shutil.copy(file_path, dest_csv)
            
            from meu_agente_cli.main import console
            with console.capture() as capture:
                handle_slash_command("/finance import")
            output = strip_ansi(capture.get()).strip()
            bot.reply_to(message, f"✅ *Resultado da Importação:*\n{output}", parse_mode="Markdown")
            return
            
        # Para outros tipos de documentos, extraímos o texto
        text_content = ""
        ext = os.path.splitext(file_name.lower())[1]
        
        if ext == ".pdf":
            text_content = extract_pdf_text(file_path)
        elif ext == ".docx":
            text_content = extract_docx_text(file_path)
        elif ext in [".xlsx", ".xls"]:
            text_content = extract_xlsx_text(file_path)
        elif ext in [".txt", ".log", ".json", ".py", ".sql", ".csv", ".xml", ".ini", ".yaml", ".yml", ".md"]:
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    text_content = f.read()
            except Exception as e:
                text_content = f"[Erro ao ler arquivo de texto: {e}]"
        else:
            bot.reply_to(message, f"❌ Formato de arquivo `{ext}` não suportado para análise automática.")
            return

        if not text_content.strip():
            bot.reply_to(message, "⚠️ O arquivo enviado parece estar vazio ou não foi possível extrair o texto dele.")
            return
            
        # Limita o tamanho do texto extraído para não estourar o limite de tokens
        if len(text_content) > 25000:
            text_content = text_content[:25000] + "\n\n... (conteúdo longo do arquivo truncado para análise) ..."
            
        caption = message.caption.strip() if message.caption else "Analise e resuma o conteúdo deste arquivo."
        user_prompt = (
            f"[Arquivo Anexado: {file_name}]\n"
            f"--- Início do Conteúdo do Arquivo ---\n"
            f"{text_content}\n"
            f"--- Fim do Conteúdo do Arquivo ---\n\n"
            f"Instrução do Usuário: {caption}"
            f"{TELEGRAM_INSTRUCTION}"
        )
        
        bot.reply_to(message, f"📖 Arquivo `{file_name}` processado ({len(text_content)} caracteres). Enviando para análise da inteligência artificial...")
        
        response = agent.process_agent_turn_silent(user_prompt)
        reply_formatted(message, response)
        
    except Exception as e:
        bot.reply_to(message, f"❌ Ocorreu um erro ao processar o arquivo: {e}")

@bot.message_handler(content_types=['photo'])
def handle_photo_upload(message):
    if not is_authorized(message):
        print(f"[BLOQUEADO] Imagem recebida de Chat ID não autorizado: {message.chat.id}")
        bot.reply_to(message, "Acesso negado. Este bot do agente é privado.")
        return

    if not is_logged_in(message):
        return

    bot.send_chat_action(message.chat.id, 'typing')
    
    try:
        import base64
        # Pega a foto na maior resolução disponível
        photo = message.photo[-1]
        file_info = bot.get_file(photo.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        
        # Salva o arquivo físico em uploads/archive
        archive_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads", "archive"))
        os.makedirs(archive_dir, exist_ok=True)
        timestamp = int(time.time())
        filename = f"image_{timestamp}.jpg"
        image_path = os.path.join(archive_dir, filename)
        
        with open(image_path, "wb") as f_img:
            f_img.write(downloaded_file)
            
        # Codifica a imagem em base64
        base64_str = base64.b64encode(downloaded_file).decode('utf-8')
        
        # Legenda/pergunta do usuário
        caption = message.caption.strip() if message.caption else "O que está nesta imagem? Descreva e analise em detalhes."
        
        # Informa o caminho da imagem no prompt para que a LLM o utilize em ferramentas se necessário
        image_info = f"\n\n[IMAGEM SALVA NO SERVIDOR: O arquivo físico desta imagem foi salvo no caminho: {image_path}]"
        
        # Monta a estrutura da mensagem multimodal
        multimodal_prompt = [
            {"type": "text", "text": f"Questão:{caption}\n\nInstruções adicionais: não utilize nenhuma ferramenta, apenas responda a questão."},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_str}"}}
        ]
        
        bot.reply_to(message, "📸 Imagem recebida e codificada. Enviando para análise visual da inteligência artificial...")
        
        response = agent.process_agent_turn_silent(multimodal_prompt, use_sys_prompt=False, use_history=False)
        reply_formatted(message, response)
        
    except Exception as e:
        logging.exception(f"Erro ao processar a imagem recebida: {e}")
        bot.reply_to(message, f"❌ Ocorreu um erro ao processar a imagem: {e}")

if __name__ == "__main__":
    # Inicializa banco de dados, LLM e segurança do projeto apenas quando executado diretamente
    if not initialize_components():
        logging.critical("Falha ao inicializar componentes críticos do agente.")
        print("[ERRO] Falha ao inicializar componentes críticos do agente.")
        sys.exit(1)

    logging.info("Bot do Telegram do Meu Agente CLI Iniciado com Sucesso!")
    print(f"===========================================================")
    print(f"Bot do Telegram do Meu Agente CLI Iniciado com Sucesso!")
    print(f"IDs Autorizados cadastrados: {authorized_ids}")
    print(f"Aguardando interações...")
    print(f"===========================================================")
    bot.infinity_polling()
